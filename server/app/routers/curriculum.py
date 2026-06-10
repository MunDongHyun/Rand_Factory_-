from datetime import datetime, timezone, timedelta
import logging
from urllib.parse import quote
from fastapi import APIRouter, Depends, HTTPException, Query as QueryParam, Request, status, Body
from fastapi.responses import StreamingResponse
from sqlalchemy import func, text
from sqlalchemy.orm import Query, Session
import os, io, json
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

from app.core.database import get_db
from app.core.db_helpers import fetch_in_order
from app.core.limiter import limiter
from app.core.security import get_current_user
from app.models.certificate import Certificate
from app.models.curriculum import Curriculum
from app.models.task_submission import TaskSubmission
from app.models.user import User
from app.services.attachment_storage import save_object, AttachmentStorageError

logger = logging.getLogger(__name__)


def _save_curriculum_export(
    curriculum_data: dict, fmt: str, content: bytes, content_type: str
) -> None:
    """매니저가 다운로드한 커리큘럼 산출물을 R2(또는 로컬)에 best-effort 보관.

    실패해도 다운로드 자체에는 영향을 주지 않고 로그만 남김.
    cur_id가 없으면 식별 불가라 저장 건너뜀.
    """
    cur_id = curriculum_data.get("cur_id")
    if not cur_id:
        return
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S")
    key = f"curriculum_exports/{cur_id}/{ts}.{fmt}"
    try:
        save_object(key, content, content_type)
    except AttachmentStorageError as exc:
        logger.warning("Curriculum export save failed cur_id=%s fmt=%s: %s", cur_id, fmt, exc)
from app.schemas.curriculum import (
    CurriculumCreate,
    CurriculumGenerateRequest,
    CurriculumGenerateResponse,
    CurriculumResponse,
    CurriculumStatsResponse,
    CurriculumUpdate,
    TemplateGenerateRequest,
    TemplateGenerateResponse,
)
from app.services import curriculum_service

router = APIRouter(prefix="/api/curricula", tags=["curricula"])


def _weeks_with_template_changed(
    old_wp: object,
    new_wp: object,
) -> set[int]:
    """양식(`template_content`)이 기존 set 상태에서 다른 값으로 바뀐 주차 집합.

    첫 set(기존이 비어 있다가 새 값 입력)은 변경으로 보지 않음 → 잠금 대상 아님.
    """
    if not isinstance(old_wp, list) or not isinstance(new_wp, list):
        return set()
    old_by_week: dict[int, list] = {}
    for w in old_wp:
        if isinstance(w, dict) and w.get("week") is not None:
            old_by_week[w["week"]] = w.get("assignments") or []
    changed: set[int] = set()
    for w in new_wp:
        if not isinstance(w, dict):
            continue
        week_no = w.get("week")
        if week_no is None:
            continue
        new_as = w.get("assignments") or []
        old_as = old_by_week.get(week_no, [])
        for i, na in enumerate(new_as):
            if not isinstance(na, dict):
                continue
            new_tc = (na.get("template_content") or "").strip()
            oa = old_as[i] if i < len(old_as) and isinstance(old_as[i], dict) else {}
            old_tc = (oa.get("template_content") or "").strip()
            if old_tc and new_tc != old_tc:
                changed.add(week_no)
                break
    return changed


def _validate_assigned_learners(
    learner_ids: list[int] | None,
    user: User,
    db: Session,
) -> list[int] | None:
    if learner_ids is None:
        return None

    unique_ids = list(dict.fromkeys(learner_ids))
    if not unique_ids:
        return []

    query = db.query(User).filter(
        User.user_id.in_(unique_ids),
        User.user_role == "j",
        User.user_deleted_at.is_(None),
    )
    if user.user_role == "m":
        query = query.filter(User.user_company == user.user_company)

    valid_ids = {learner.user_id for learner in query.all()}
    invalid_ids = [learner_id for learner_id in unique_ids if learner_id not in valid_ids]
    if invalid_ids:
        raise HTTPException(status_code=400, detail="배정할 수 없는 학습자가 포함되어 있습니다.")

    return unique_ids


def _scope_curriculum_query(query: Query, user: User) -> Query:
    query = query.filter(Curriculum.cur_deleted_at.is_(None))
    if user.user_role == "a":
        return query
    if user.user_role == "m":
        return query.filter(Curriculum.cur_creator_id == user.user_id)
    if user.user_role == "j":
        # JSON_CONTAINS 두 번째 인자는 JSON 문서여야 하므로 json.dumps 로 명시
        return query.filter(
            func.json_contains(
                Curriculum.cur_assigned_learner_ids,
                json.dumps(user.user_id),
            )
        )
    return query.filter(False)


@router.post("", response_model=CurriculumResponse, status_code=status.HTTP_201_CREATED)
def create_curriculum(
    body: CurriculumCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=403, detail="커리큘럼은 매니저/관리자만 생성할 수 있습니다")

    curriculum = Curriculum(
        cur_creator_id=current_user.user_id,
        cur_title=body.cur_title,
        cur_target_job=body.cur_target_job,
        cur_target_industry=body.cur_target_industry,
        cur_duration_weeks=body.cur_duration_weeks,
        cur_learning_goal=body.cur_learning_goal,
        cur_learning_detail_goal=body.cur_learning_detail_goal,
        cur_week_plan=body.cur_week_plan,
        cur_assigned_learner_ids=_validate_assigned_learners(body.cur_assigned_learner_ids, current_user, db),
        cur_status=body.cur_status,
    )
    db.add(curriculum)
    db.commit()
    db.refresh(curriculum)
    return curriculum


@router.post("/generate", response_model=CurriculumGenerateResponse)
@limiter.limit("10/hour")
def generate_curriculum(
    request: Request,
    body: CurriculumGenerateRequest,
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=403, detail="커리큘럼 자동 생성은 매니저/관리자만 사용할 수 있습니다")

    try:
        val_result = curriculum_service.validate_curriculum_input(
            cur_title=body.cur_title,
            cur_target_job=body.cur_target_job or "",
            cur_target_industry=body.cur_target_industry or "",
            cur_learning_goal=body.cur_learning_goal or "",
            required_content=body.required_content or ""
        )
        if not val_result.is_valid:
            raise HTTPException(status_code=400, detail=val_result.reason)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"입력값 검증 실패: {exc}")

    try:
        week_plan = curriculum_service.generate_week_plan(
            cur_title=body.cur_title,
            cur_duration_weeks=body.cur_duration_weeks,
            cur_target_job=body.cur_target_job,
            cur_target_industry=body.cur_target_industry,
            cur_learning_goal=body.cur_learning_goal,
            required_content=body.required_content,
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"AI 생성 실패: {exc}")

    return CurriculumGenerateResponse(
        cur_title=body.cur_title,
        cur_duration_weeks=body.cur_duration_weeks,
        cur_target_job=body.cur_target_job,
        cur_target_industry=body.cur_target_industry,
        cur_learning_goal=body.cur_learning_goal,
        cur_week_plan=week_plan,
    )


@router.get("", response_model=list[CurriculumResponse])
def list_curricula(
    limit: int = QueryParam(100, ge=1, le=500),
    offset: int = QueryParam(0, ge=0),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # sort buffer 폭발 방지 + 페이지네이션
    id_query = (
        _scope_curriculum_query(db.query(Curriculum.cur_id), current_user)
        .order_by(Curriculum.cur_created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    return fetch_in_order(db, id_query, Curriculum, Curriculum.cur_id)


@router.get("/stats", response_model=CurriculumStatsResponse)
def get_curriculum_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role != "a":
        raise HTTPException(status_code=404, detail="찾을 수 없습니다")

    total_curricula = (
        db.query(func.count(Curriculum.cur_id))
        .filter(Curriculum.cur_deleted_at.is_(None))
        .scalar()
        or 0
    )

    # MySQL 8 JSON_TABLE 로 활성 커리큘럼의 배정 학습자 합집합 cardinality 를 DB 측에서 계산.
    # NULL 컬럼은 IFNULL 로 빈 배열 처리해 안전.
    active_learners = int(
        db.execute(
            text(
                """
                SELECT COUNT(DISTINCT jt.learner_id)
                FROM curriculum c
                CROSS JOIN JSON_TABLE(
                    IFNULL(c.cur_assigned_learner_ids, JSON_ARRAY()),
                    '$[*]' COLUMNS(learner_id BIGINT PATH '$')
                ) jt
                WHERE c.cur_deleted_at IS NULL AND c.cur_status = 'active'
                """
            )
        ).scalar()
        or 0
    )

    total_submissions = (
        db.query(func.count(TaskSubmission.task_submission_id)).scalar() or 0
    )

    return CurriculumStatsResponse(
        total_curricula=total_curricula,
        active_learners=active_learners,
        total_submissions=total_submissions,
    )


@router.get("/{cur_id}/completion-report")
def download_completion_report(
    cur_id: int,
    comment: str | None = QueryParam(None, max_length=1000),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    curriculum = _load_curriculum_for_manager_report(db, current_user, cur_id)
    expected_weeks = _expected_week_numbers(curriculum)
    learner_ids = _assigned_ids(curriculum)

    learners: list[User] = []
    if learner_ids:
        learner_rows = (
            db.query(User)
            .filter(
                User.user_id.in_(learner_ids),
                User.user_role == "j",
                User.user_deleted_at.is_(None),
            )
            .all()
        )
        learner_map = {learner.user_id: learner for learner in learner_rows}
        learners = [learner_map[learner_id] for learner_id in learner_ids if learner_id in learner_map]

    submissions = _latest_report_submissions(db, curriculum.cur_id)
    latest_by_learner_week: dict[tuple[int, int], TaskSubmission] = {}
    learner_id_set = {learner.user_id for learner in learners}
    expected_set = set(expected_weeks)
    for submission in submissions:
        if submission.task_learner_id not in learner_id_set:
            continue
        if submission.task_week_number not in expected_set:
            continue
        latest_by_learner_week[(submission.task_learner_id, submission.task_week_number)] = submission

    cert_rows = (
        db.query(Certificate)
        .filter(
            Certificate.cert_curriculum_id == curriculum.cur_id,
            Certificate.cert_deleted_at.is_(None),
        )
        .all()
    )
    certified_ids = {
        cert.cert_learner_id
        for cert in cert_rows
        if cert.cert_learner_id in learner_id_set
    }

    creator = db.query(User).filter(User.user_id == curriculum.cur_creator_id).first()
    generated_at = datetime.now(timezone.utc)
    pdf_bytes = _generate_completion_report_pdf(
        curriculum=curriculum,
        creator=creator,
        learners=learners,
        expected_weeks=expected_weeks,
        latest_by_learner_week=latest_by_learner_week,
        certified_ids=certified_ids,
        manager_comment=comment,
        generated_at=generated_at,
    )

    key = f"curriculum_reports/{curriculum.cur_id}/{generated_at.strftime('%Y%m%dT%H%M%S')}.pdf"
    try:
        save_object(key, pdf_bytes, "application/pdf")
    except AttachmentStorageError as exc:
        logger.warning("Completion report save failed cur_id=%s: %s", curriculum.cur_id, exc)

    filename = f"{curriculum.cur_title or 'curriculum'}_completion_report.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers=_download_headers(filename),
    )


@router.get("/{cur_id}", response_model=CurriculumResponse)
def get_curriculum(
    cur_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = _scope_curriculum_query(db.query(Curriculum), current_user)
    curriculum = query.filter(Curriculum.cur_id == cur_id).first()
    if not curriculum:
        raise HTTPException(status_code=404, detail="커리큘럼을 찾을 수 없습니다")
    return curriculum


@router.patch("/{cur_id}", response_model=CurriculumResponse)
def update_curriculum(
    cur_id: int,
    body: CurriculumUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=404, detail="커리큘럼을 찾을 수 없습니다")

    query = db.query(Curriculum).filter(
        Curriculum.cur_id == cur_id,
        Curriculum.cur_deleted_at.is_(None),
    )
    if current_user.user_role == "m":
        query = query.filter(Curriculum.cur_creator_id == current_user.user_id)

    curriculum = query.first()
    if not curriculum:
        raise HTTPException(status_code=404, detail="커리큘럼을 찾을 수 없습니다")

    update_data = body.model_dump(exclude_unset=True)
    if "cur_assigned_learner_ids" in update_data:
        update_data["cur_assigned_learner_ids"] = _validate_assigned_learners(
            update_data["cur_assigned_learner_ids"],
            current_user,
            db,
        )

    if "cur_week_plan" in update_data:
        changed_weeks = _weeks_with_template_changed(
            curriculum.cur_week_plan,
            update_data["cur_week_plan"],
        )
        if changed_weeks:
            locked = (
                db.query(TaskSubmission.task_week_number)
                .filter(
                    TaskSubmission.task_curriculum_id == cur_id,
                    TaskSubmission.task_week_number.in_(list(changed_weeks)),
                )
                .distinct()
                .all()
            )
            locked_weeks = sorted({r.task_week_number for r in locked})
            if locked_weeks:
                week_str = ", ".join(f"{w}주차" for w in locked_weeks)
                raise HTTPException(
                    status_code=400,
                    detail=f"이미 학습자 제출이 있는 양식은 수정할 수 없습니다: {week_str}",
                )

    for field, value in update_data.items():
        setattr(curriculum, field, value)

    db.commit()
    db.refresh(curriculum)
    return curriculum


@router.post("/download/txt")
async def download_curriculum_txt(
    curriculum_data: dict = Body(...),
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=404, detail="찾을 수 없습니다")
    
    cur_title = curriculum_data.get("cur_title", "커리큘럼")
    week_plan = curriculum_data.get("cur_week_plan", [])
    if isinstance(week_plan, str):
        week_plan = json.loads(week_plan)

    lines = []
    lines.append("=" * 60)
    lines.append(" ArtiCulum 과제 템플릿 패키지")
    lines.append(f" {cur_title} 실무 온보딩 과정")
    lines.append("=" * 60 + "\n")

    # 1. 과정 개요
    lines.append("1. 과정 개요")
    lines.append("-" * 60)
    lines.append(f"과정명: {cur_title}")
    lines.append(f"대상 직무: {curriculum_data.get('cur_target_job', '미지정')}")
    lines.append(f"교육 기간: {curriculum_data.get('cur_duration_weeks', 0)}주")
    lines.append(f"교육 목적: {curriculum_data.get('cur_learning_goal', '미지정')}")
    lines.append("-" * 60 + "\n")

    # 2. 전체 과제 흐름
    lines.append("2. 전체 과제 흐름")
    lines.append("-" * 60)
    for week in week_plan:
        tasks = ", ".join([a.get("title", "") for a in week.get("assignments", [])])
        lines.append(f"[{week.get('week')}주차] 주제: {week.get('theme')}")
        lines.append(f"        핵심 산출물: {tasks}")
    lines.append("-" * 60 + "\n")

    # 3. 주차별 과제 템플릿
    lines.append("3. 주차별 과제 템플릿")
    for week in week_plan:
        lines.append("=" * 60)
        lines.append(f" [ {week.get('week')}주차 과제 템플릿 ]")
        lines.append(f" 주제: {week.get('theme')}")
        lines.append(f" 학습 목표: {week.get('learning_objective')}")
        lines.append("=" * 60)
        
        for idx, a in enumerate(week.get("assignments", [])):
            lines.append(f"\n▶ 과제명: {a.get('title')}")
            lines.append(f"  제출 형태: {a.get('expected_output_format')}")
            lines.append(f"  수행 방법:")
            for step in a.get("step_by_step_guide", []):
                lines.append(f"    - {step}")
                
        # 자기 점검 & 사수 피드백
        ig = week.get("instructor_guide", {})
        if ig and ig.get("check_points"):
            lines.append("\n[ 자기 점검 체크리스트 ]")
            for cp in ig.get("check_points", []):
                lines.append(f"  ☐ {cp}")
                
        if ig and ig.get("coaching_questions"):
            lines.append("\n[ 사수 피드백 질문 ]")
            for cq in ig.get("coaching_questions", []):
                lines.append(f"  Q. {cq}")
                
        lines.append("\n")

    # 4. 공통 평가 루브릭
    lines.append("4. 공통 평가 루브릭")
    lines.append("-" * 60)
    lines.append("구체성: 행동, 수량, 순서, 보고 대상이 명확한가?")
    lines.append("현장 적용성: 실제 업무에서 바로 사용할 수 있는가?")
    lines.append("근거 제시: 명확한 데이터나 상황적 근거가 포함되었는가?")
    lines.append("연결성: 이전 주차 산출물을 다음 과제에 잘 활용했는가?")
    lines.append("-" * 60 + "\n")

    text_content = "\n".join(lines)
    _save_curriculum_export(curriculum_data, "txt", text_content.encode("utf-8"), "text/plain; charset=utf-8")

    return StreamingResponse(
        io.StringIO(text_content),
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename=curriculum.txt"},
    )


current_file = os.path.abspath(__file__)
router_dir = os.path.dirname(current_file)
app_dir = os.path.dirname(router_dir)
server_dir = os.path.dirname(app_dir)
FONT_PATH = os.path.join(server_dir, "resources", "fonts", "NanumMyeongjo-subset.ttf")


def _download_headers(filename: str) -> dict[str, str]:
    fallback = filename.encode("ascii", "ignore").decode("ascii") or "download.pdf"
    encoded = quote(filename, safe="")
    return {
        "Content-Disposition": f"attachment; filename=\"{fallback}\"; filename*=UTF-8''{encoded}"
    }


def _safe_pdf_text(value: object, fallback: str = "-") -> str:
    text_value = fallback if value is None else str(value)
    text_value = text_value.replace("\r", "").replace("\t", "  ").strip()
    return text_value or fallback


def _expected_week_numbers(curriculum: Curriculum) -> list[int]:
    week_plan = curriculum.cur_week_plan
    if isinstance(week_plan, list) and week_plan:
        weeks = sorted(
            {
                int(item["week"])
                for item in week_plan
                if isinstance(item, dict)
                and item.get("week") is not None
                and str(item.get("week")).isdigit()
            }
        )
        if weeks:
            return weeks
    if isinstance(week_plan, dict) and week_plan.get("week") is not None:
        raw_week = week_plan.get("week")
        if str(raw_week).isdigit():
            return [int(raw_week)]
    return list(range(1, int(curriculum.cur_duration_weeks or 0) + 1))


def _assigned_ids(curriculum: Curriculum) -> list[int]:
    raw = curriculum.cur_assigned_learner_ids
    if not isinstance(raw, list):
        return []
    seen: set[int] = set()
    ids: list[int] = []
    for value in raw:
        if isinstance(value, int) or str(value).isdigit():
            learner_id = int(value)
            if learner_id not in seen:
                ids.append(learner_id)
                seen.add(learner_id)
    return ids


def _load_curriculum_for_manager_report(
    db: Session,
    current_user: User,
    cur_id: int,
) -> Curriculum:
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=404, detail="Not found")
    query = db.query(Curriculum).filter(
        Curriculum.cur_id == cur_id,
        Curriculum.cur_deleted_at.is_(None),
    )
    if current_user.user_role == "m":
        query = query.filter(Curriculum.cur_creator_id == current_user.user_id)
    curriculum = query.first()
    if not curriculum:
        raise HTTPException(status_code=404, detail="Curriculum not found")
    return curriculum


def _latest_report_submissions(
    db: Session,
    curriculum_id: int,
) -> list[TaskSubmission]:
    id_query = (
        db.query(TaskSubmission.task_submission_id)
        .filter(TaskSubmission.task_curriculum_id == curriculum_id)
        .order_by(
            TaskSubmission.task_learner_id.asc(),
            TaskSubmission.task_week_number.asc(),
            TaskSubmission.task_submitted_at.asc(),
            TaskSubmission.task_submission_id.asc(),
        )
    )
    return fetch_in_order(db, id_query, TaskSubmission, TaskSubmission.task_submission_id)


def _generate_completion_report_pdf(
    *,
    curriculum: Curriculum,
    creator: User | None,
    learners: list[User],
    expected_weeks: list[int],
    latest_by_learner_week: dict[tuple[int, int], TaskSubmission],
    certified_ids: set[int],
    manager_comment: str | None,
    generated_at: datetime,
) -> bytes:
    # 보고서 정체성: 매니저가 윗선(임원/HR)에 제출하는 교육 결과 보고서.
    # - 총평(매니저 코멘트)이 데이터보다 앞쪽
    # - 수료자 명단을 별도 박스로 강조
    # - 학습자 표는 성과 좋은 순서로 정렬 (발급 완료 → 발급 가능 → 진행 중)
    # - 개인정보 노출 최소화 (이메일 제거)
    # - 페이지 푸터에 페이지 번호 + 담당 매니저

    has_font = os.path.exists(FONT_PATH)
    issuer_name = creator.user_name if creator else "-"

    class _ReportPDF(FPDF):
        def footer(self):
            # 매 페이지 푸터: 담당 매니저 + 페이지 번호. 페이지 번호 표현은 alias_nb_pages()로 총 페이지 치환.
            self.set_y(-12)
            self.set_font("NanumGothic" if has_font else "Arial", "", 8)
            self.set_text_color(150, 156, 165)
            self.cell(95, 5, _safe_pdf_text(f"담당 매니저: {issuer_name}"), align="L")
            self.cell(95, 5, f"{self.page_no()} / {{nb}}", align="R")

    pdf = _ReportPDF(orientation="P", unit="mm", format="A4")
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=False)
    if has_font:
        pdf.add_font("NanumGothic", "", FONT_PATH, uni=True)
    pdf.add_page()

    def set_font(size: int = 10):
        pdf.set_font("NanumGothic" if has_font else "Arial", "", size)

    def ensure_space(height: float):
        # 푸터 영역(약 15mm) 침범 방지
        if pdf.get_y() + height > 275:
            pdf.add_page()

    def percent(part: int, whole: int) -> int:
        return int(round(100 * part / whole)) if whole else 0

    def text_line(label: str, value: object):
        ensure_space(8)
        set_font(9)
        pdf.set_text_color(90, 96, 108)
        pdf.cell(33, 7, _safe_pdf_text(label), ln=False)
        pdf.set_text_color(28, 35, 45)
        pdf.multi_cell(0, 7, _safe_pdf_text(value))

    def section(title: str):
        ensure_space(14)
        pdf.ln(3)
        set_font(12)
        pdf.set_text_color(31, 41, 55)
        pdf.cell(0, 8, title, ln=True)
        pdf.set_draw_color(220, 225, 232)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)

    def summary_card(x: float, y: float, width: float, label: str, value: str, subtext: str):
        pdf.set_draw_color(220, 225, 232)
        pdf.set_fill_color(248, 250, 252)
        pdf.rect(x, y, width, 23, "DF")
        pdf.set_xy(x + 4, y + 3)
        set_font(8)
        pdf.set_text_color(90, 96, 108)
        pdf.cell(width - 8, 5, label, ln=True)
        pdf.set_xy(x + 4, y + 9)
        set_font(14)
        pdf.set_text_color(31, 41, 55)
        pdf.cell(width - 8, 7, value, ln=True)
        pdf.set_xy(x + 4, y + 17)
        set_font(7)
        pdf.set_text_color(110, 118, 130)
        pdf.cell(width - 8, 4, subtext)

    learner_ids = [learner.user_id for learner in learners]
    total_slots = len(learner_ids) * len(expected_weeks)
    completed_slots = 0
    completed_learners = 0
    learner_rows: list[dict[str, object]] = []
    week_rows: list[dict[str, object]] = []

    def _aware(dt):
        # task_submitted_at / task_deadline 비교용. naive datetime은 UTC로 가정.
        if dt is None:
            return None
        if dt.tzinfo is None:
            return dt.replace(tzinfo=timezone.utc)
        return dt

    for learner in learners:
        submitted_weeks = []
        feedback_weeks = []
        resubmit_weeks = []
        pending_weeks = []
        missing_weeks = []
        for week in expected_weeks:
            submission = latest_by_learner_week.get((learner.user_id, week))
            if not submission:
                missing_weeks.append(week)
                continue
            submitted_weeks.append(week)
            if submission.task_resubmit_requested == "Y" or submission.task_status == "resubmit_requested":
                resubmit_weeks.append(week)
            elif submission.task_status == "feedback_given":
                feedback_weeks.append(week)
            else:
                pending_weeks.append(week)
        learner_complete = (
            bool(expected_weeks)
            and len(feedback_weeks) == len(expected_weeks)
            and not resubmit_weeks
        )
        completed_slots += len(feedback_weeks)
        if learner_complete:
            completed_learners += 1
        # 수료 상태 3단계: 발급 완료 / 발급 가능 / 진행 중
        is_certified = learner.user_id in certified_ids
        if learner_complete and is_certified:
            cert_status = "발급 완료"
            sort_rank = 0
        elif learner_complete:
            cert_status = "발급 가능"
            sort_rank = 1
        else:
            cert_status = "진행 중"
            sort_rank = 2
        progress_for_learner = percent(len(feedback_weeks), len(expected_weeks))
        learner_rows.append({
            "name": learner.user_name,
            "submitted": len(set(submitted_weeks)),
            "feedback": len(set(feedback_weeks)),
            "resubmit": len(set(resubmit_weeks)),
            "certified": is_certified,
            "complete": learner_complete,
            "cert_status": cert_status,
            "sort_rank": sort_rank,
            "progress_pct": progress_for_learner,
            "missing_weeks": sorted(set(missing_weeks)),
            "pending_weeks": sorted(set(pending_weeks)),
            "resubmit_weeks": sorted(set(resubmit_weeks)),
        })

    # 학습자 표 정렬: 발급 완료 → 발급 가능 → 진행 중. 동순위는 진행률 desc, 이름 asc.
    learner_rows.sort(key=lambda r: (r["sort_rank"], -int(r["progress_pct"]), r["name"] or ""))

    # 마감일 준수 현황 — 주차별 정시·지연 통계
    total_with_deadline = 0
    total_on_time = 0
    total_late = 0
    for week in expected_weeks:
        submitted = 0
        feedback = 0
        on_time = 0
        late = 0
        with_deadline = 0
        missing_names = []
        for learner in learners:
            submission = latest_by_learner_week.get((learner.user_id, week))
            if submission:
                submitted += 1
                if (
                    submission.task_status == "feedback_given"
                    and submission.task_resubmit_requested != "Y"
                ):
                    feedback += 1
                deadline = _aware(submission.task_deadline)
                submitted_at = _aware(submission.task_submitted_at)
                if deadline and submitted_at:
                    with_deadline += 1
                    if submitted_at <= deadline:
                        on_time += 1
                    else:
                        late += 1
            else:
                missing_names.append(learner.user_name)
        total_with_deadline += with_deadline
        total_on_time += on_time
        total_late += late
        week_rows.append({
            "week": week,
            "submitted": submitted,
            "feedback": feedback,
            "missing": missing_names,
            "on_time": on_time,
            "late": late,
            "with_deadline": with_deadline,
        })

    progress_pct = percent(completed_slots, total_slots)
    completion_pct = percent(completed_learners, len(learner_ids))
    certified_total = len(certified_ids)
    overall_on_time_pct = percent(total_on_time, total_with_deadline) if total_with_deadline else None

    set_font(18)
    pdf.set_text_color(22, 28, 38)
    pdf.cell(0, 12, "ArtiCulum 커리큘럼 완료보고서", ln=True, align="C")
    set_font(9)
    pdf.set_text_color(90, 96, 108)
    pdf.cell(0, 6, f"생성일: {generated_at.strftime('%Y-%m-%d %H:%M')}", ln=True, align="C")
    pdf.ln(5)

    section("1. 커리큘럼 개요")
    text_line("과정명", curriculum.cur_title)
    text_line("교육 대상", curriculum.cur_target_job or "-")
    text_line("적용 범위", curriculum.cur_target_industry or "-")
    text_line("기간", f"{curriculum.cur_duration_weeks or len(expected_weeks)}주")
    text_line("학습 목표", curriculum.cur_learning_goal or "-")
    text_line("담당 매니저", issuer_name)

    section("2. 운영 요약")
    # 성과 인사이트 한 줄 — 카드 위 자연어 요약. 결과 보고서 첫인상 강화.
    insight_parts = [
        f"배정 학습자 {len(learner_ids)}명 중 {completed_learners}명 수료 (수료율 {completion_pct}%)",
        f"전체 진행률 {progress_pct}%",
    ]
    if overall_on_time_pct is not None:
        insight_parts.append(f"마감 준수율 {overall_on_time_pct}%")
    set_font(9)
    pdf.set_text_color(31, 41, 55)
    pdf.multi_cell(0, 6, ". ".join(insight_parts) + ".")
    pdf.ln(1)

    card_y = pdf.get_y() + 1
    card_w = 43.5
    card_gap = 4
    start_x = 10
    summary_card(start_x, card_y, card_w, "배정 학습자", f"{len(learner_ids)}명", "보고서 대상")
    summary_card(start_x + (card_w + card_gap), card_y, card_w, "전체 진행률", f"{progress_pct}%", f"{completed_slots}/{total_slots} 피드백 완료")
    summary_card(start_x + (card_w + card_gap) * 2, card_y, card_w, "수료율", f"{completion_pct}%", f"{completed_learners}/{len(learner_ids)}명 수료")
    summary_card(start_x + (card_w + card_gap) * 3, card_y, card_w, "수료증 발급", f"{certified_total}명", "발급 완료")
    pdf.set_y(card_y + 28)
    set_font(8)
    pdf.set_text_color(90, 96, 108)
    pdf.multi_cell(0, 6, "완료 기준: 모든 예정 주차에 제출이 있고, 재제출 요청 없이 매니저 피드백이 완료된 상태입니다.")
    pdf.set_text_color(28, 35, 45)

    # 매니저 종합 의견 — 결과 보고서 흐름상 데이터보다 앞쪽에 배치
    next_section_no = 3
    if manager_comment:
        section(f"{next_section_no}. 매니저 종합 의견")
        set_font(9)
        pdf.set_text_color(31, 41, 55)
        pdf.multi_cell(0, 7, _safe_pdf_text(manager_comment))
        next_section_no += 1

    # 수료자 명단 박스 — 윗선 보고에서 가장 중요한 정보
    section(f"{next_section_no}. 수료자 명단")
    next_section_no += 1
    if certified_total == 0:
        set_font(9)
        pdf.set_text_color(110, 118, 130)
        pdf.multi_cell(0, 7, "이번 보고 시점 기준 수료증 발급 완료자가 없습니다.")
        pdf.set_text_color(28, 35, 45)
    else:
        certified_names = [row["name"] for row in learner_rows if row["certified"]]
        ensure_space(20)
        box_x = 10
        box_y = pdf.get_y()
        # 이름이 많아질 경우 박스 높이 동적 계산
        approx_line_count = max(1, ((sum(len(n) for n in certified_names) + len(certified_names) * 2) // 70) + 1)
        box_h = 8 + approx_line_count * 6
        pdf.set_draw_color(199, 213, 175)
        pdf.set_fill_color(244, 249, 235)
        pdf.rect(box_x, box_y, 190, box_h, "DF")
        pdf.set_xy(box_x + 4, box_y + 3)
        set_font(8)
        pdf.set_text_color(90, 110, 60)
        pdf.cell(0, 5, f"수료증 발급 완료 {certified_total}명", ln=True)
        pdf.set_xy(box_x + 4, box_y + 9)
        set_font(10)
        pdf.set_text_color(31, 41, 55)
        pdf.multi_cell(190 - 8, 6, _safe_pdf_text(", ".join(certified_names)))
        pdf.set_y(box_y + box_h + 2)
        pdf.set_text_color(28, 35, 45)

    # 의미 단위 분할: 1페이지는 요약/총평/수료자(윗선 1분 요약), 2페이지는 상세 데이터.
    # 학습자가 많아 자연히 페이지가 늘어나는 케이스에서도 이 구분은 그대로 유지된다.
    pdf.add_page()

    section(f"{next_section_no}. 학습자별 성과")
    next_section_no += 1
    set_font(8)
    pdf.set_fill_color(244, 246, 249)
    pdf.set_text_color(31, 41, 55)
    headers = [
        ("학습자", 42),
        ("제출률", 28),
        ("진행률", 22),
        ("피드백", 22),
        ("재제출", 22),
    ]
    for text, width in headers:
        pdf.cell(width, 7, text, border=1, fill=True, align="C")
    pdf.cell(0, 7, "수료 상태", border=1, fill=True, align="C", ln=True)
    for row in learner_rows:
        ensure_space(8)
        pdf.cell(42, 7, _safe_pdf_text(row["name"]), border=1)
        pdf.cell(28, 7, f"{row['submitted']}/{len(expected_weeks)}", border=1, align="C")
        pdf.cell(22, 7, f"{row['progress_pct']}%", border=1, align="C")
        pdf.cell(22, 7, str(row["feedback"]), border=1, align="C")
        pdf.cell(22, 7, str(row["resubmit"]), border=1, align="C")
        pdf.cell(0, 7, _safe_pdf_text(row["cert_status"]), border=1, align="C", ln=True)

    section(f"{next_section_no}. 주차별 제출 현황")
    next_section_no += 1
    set_font(8)
    pdf.set_fill_color(244, 246, 249)
    pdf.set_text_color(31, 41, 55)
    week_headers = [
        ("주차", 18),
        ("제출률", 32),
        ("피드백 완료율", 32),
        ("정시 / 지연", 32),
    ]
    for text, width in week_headers:
        pdf.cell(width, 7, text, border=1, fill=True, align="C")
    pdf.cell(0, 7, "미제출자", border=1, fill=True, align="C", ln=True)
    for row in week_rows:
        ensure_space(8)
        rate = percent(int(row["submitted"]), len(learner_ids))
        feedback_rate = percent(int(row["feedback"]), len(learner_ids))
        if int(row["with_deadline"]) > 0:
            on_time_text = f"{row['on_time']} / {row['late']}"
        else:
            on_time_text = "-"
        missing = ", ".join(row["missing"][:4])
        if len(row["missing"]) > 4:
            missing += f" 외 {len(row['missing']) - 4}명"
        pdf.cell(18, 7, f"{row['week']}주차", border=1, align="C")
        pdf.cell(32, 7, f"{rate}% ({row['submitted']}/{len(learner_ids)})", border=1, align="C")
        pdf.cell(32, 7, f"{feedback_rate}%", border=1, align="C")
        pdf.cell(32, 7, on_time_text, border=1, align="C")
        pdf.cell(0, 7, _safe_pdf_text(missing or "-"), border=1, ln=True)

    if total_with_deadline > 0:
        ensure_space(8)
        pdf.ln(1)
        set_font(8)
        pdf.set_text_color(90, 96, 108)
        pdf.multi_cell(
            0,
            6,
            f"전체 마감 준수율 {overall_on_time_pct}% — 정시 {total_on_time}건 / 지연 {total_late}건 (마감 설정 {total_with_deadline}건 기준).",
        )
        pdf.set_text_color(28, 35, 45)

    # 미진행 현황 — 윗선 보고 톤: 개인 명단보다 집계 위주
    section(f"{next_section_no}. 미진행 현황")
    next_section_no += 1
    incomplete_rows = [row for row in learner_rows if not row["complete"]]
    if not incomplete_rows:
        set_font(9)
        pdf.set_text_color(31, 41, 55)
        pdf.multi_cell(0, 7, "모든 학습자가 학습 완료 기준을 충족했습니다.")
    else:
        avg_progress = (
            int(round(sum(int(r["progress_pct"]) for r in incomplete_rows) / len(incomplete_rows)))
            if incomplete_rows
            else 0
        )
        # 주차별 미진행자 수 집계 (어느 주차에서 막혔는지)
        missing_by_week: dict[int, int] = {}
        pending_by_week: dict[int, int] = {}
        resubmit_by_week: dict[int, int] = {}
        for row in incomplete_rows:
            for week in row["missing_weeks"]:
                missing_by_week[week] = missing_by_week.get(week, 0) + 1
            for week in row["pending_weeks"]:
                pending_by_week[week] = pending_by_week.get(week, 0) + 1
            for week in row["resubmit_weeks"]:
                resubmit_by_week[week] = resubmit_by_week.get(week, 0) + 1

        def _fmt_weeks(by_week: dict[int, int]) -> str:
            if not by_week:
                return "-"
            return ", ".join(f"{w}주차({n}명)" for w, n in sorted(by_week.items()))

        set_font(9)
        pdf.set_text_color(31, 41, 55)
        pdf.multi_cell(
            0,
            7,
            f"진행 중 {len(incomplete_rows)}명 / 평균 진행률 {avg_progress}%.",
        )
        pdf.ln(1)
        set_font(9)
        pdf.set_text_color(90, 96, 108)
        text_line("미제출 주차", _fmt_weeks(missing_by_week))
        text_line("피드백 대기 주차", _fmt_weeks(pending_by_week))
        text_line("재제출 요청 주차", _fmt_weeks(resubmit_by_week))
        pdf.set_text_color(28, 35, 45)

    ensure_space(12)
    pdf.ln(3)
    set_font(8)
    pdf.set_text_color(110, 118, 130)
    pdf.multi_cell(0, 5, "본 보고서는 ArtiCulum에서 자동 생성되었으며, 생성된 PDF는 객체 저장소에 보관됩니다.")

    out = pdf.output(dest="S")
    return out.encode("latin-1") if isinstance(out, str) else bytes(out)


@router.post("/download/pdf")
async def download_curriculum_pdf(
    curriculum_data: dict = Body(...),
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=404, detail="찾을 수 없습니다")
    
    cur_title = curriculum_data.get("cur_title", "커리큘럼")
    week_plan = curriculum_data.get("cur_week_plan", [])
    if isinstance(week_plan, str):
        week_plan = json.loads(week_plan)

    pdf = FPDF()
    pdf.add_page()
    
    has_font = os.path.exists(FONT_PATH)
    if has_font:
        pdf.add_font('NanumGothic', '', FONT_PATH, uni=True)
    
    def set_font(style='', size=10):
        if has_font:
            pdf.set_font('NanumGothic', style, size)
        else:
            pdf.set_font('Arial', style, size)
            
    def safe_print(text, border=0, fill=False, h=7):
        if not text: return
        text = str(text).replace('\r', '').replace('\t', '  ')
        
        max_w = pdf.w - pdf.l_margin - pdf.r_margin - 2 
        
        lines = []
        for p in text.split('\n'):
            if not p:
                lines.append("")
                continue
                
            current_line = ""
            for char in p:
                if pdf.get_string_width(current_line + char) > max_w:
                    lines.append(current_line)
                    current_line = char
                else:
                    current_line += char
            
            if current_line:
                lines.append(current_line)
        
        for i, line in enumerate(lines):
            b = 0
            if border:
                if len(lines) == 1: b = 1
                elif i == 0: b = 'LTR'
                elif i == len(lines) - 1: b = 'LBR'
                else: b = 'LR'
            pdf.cell(0, h, txt=line, border=b, ln=True, fill=fill)

    set_font('', 16)
    pdf.cell(0, 10, txt="ArtiCulum 과제 템플릿 패키지", ln=True, align='C')
    set_font('', 14)
    pdf.cell(0, 10, txt=f"{cur_title} 실무 온보딩 과정", ln=True, align='C')
    pdf.ln(5)

    # 1. 과정 개요
    set_font('', 12)
    pdf.cell(0, 10, txt="1. 과정 개요", ln=True)
    set_font('', 10)
    
    overview_text = f"과정명: {cur_title}\n대상 직무: {curriculum_data.get('cur_target_job', '미지정')}\n교육 기간: {curriculum_data.get('cur_duration_weeks', 0)}주\n교육 목적: {curriculum_data.get('cur_learning_goal', '미지정')}"
    safe_print(overview_text, border=1, h=8)
    pdf.ln(5)

    # 2. 전체 과제 흐름
    set_font('', 12)
    pdf.cell(0, 10, txt="2. 전체 과제 흐름", ln=True)
    set_font('', 10)
    
    for week in week_plan:
        week_num = week.get('week', '')
        theme = str(week.get('theme', ''))
        tasks = ", ".join([a.get("title", "") for a in week.get("assignments", [])])
        pdf.cell(0, 8, txt=f"[{week_num}주차] {theme}", ln=True)
        safe_print(f"산출물: {tasks}")

    pdf.ln(5)
    
    # 3. 주차별 상세 과제 가이드
    set_font('', 12)
    pdf.cell(0, 10, txt="3. 주차별 과제 템플릿", ln=True)
    set_font('', 10)
    
    for week in week_plan:
        pdf.set_fill_color(240, 240, 240)
        title_str = f"■ {week.get('week')}주차: {week.get('theme')}"
        pdf.cell(0, 8, txt=title_str, ln=True, fill=True)
        
        safe_print(f"학습 목표: {week.get('learning_objective')}")
        
        assignments = week.get('assignments', [])
        for idx, a in enumerate(assignments):
            pdf.ln(2)
            task_name = f"▶ 과제명: {a.get('title')}"
            pdf.cell(0, 7, txt=task_name, ln=True)
            
            safe_print(f"제출 형태: {a.get('expected_output_format')}")

            guide_text = "수행 방법:\n" + "\n".join([f"  - {g}" for g in a.get('step_by_step_guide', [])])
            safe_print(guide_text)
            
        ig = week.get('instructor_guide', {})
        if ig and ig.get('check_points'):
            pdf.ln(2)
            pdf.cell(0, 7, txt="[ 자기 점검 체크리스트 ]", ln=True)
            for cp in ig.get('check_points', []):
                safe_print(f"  ☐ {cp}")
                
        if ig and ig.get('coaching_questions'):
            pdf.ln(2)
            pdf.cell(0, 7, txt="[ 사수 피드백 질문 ]", ln=True)
            for cq in ig.get('coaching_questions', []):
                safe_print(f"  Q. {cq}")

        pdf.ln(5)

    # 4. 평가 루브릭
    set_font('', 12)
    pdf.cell(0, 10, txt="4. 공통 평가 루브릭", ln=True)
    set_font('', 10)
    rubric = "구체성: 행동/수량/순서 명확\n현장 적용성: 실무 즉시 적용 가능\n근거 제시: 데이터/상황 근거 포함\n연결성: 이전 주차 산출물 연계"
    safe_print(rubric, border=1, h=8)
    pdf_out = pdf.output(dest="S")
    pdf_bytes = pdf_out.encode("latin-1") if isinstance(pdf_out, str) else bytes(pdf_out)
    _save_curriculum_export(curriculum_data, "pdf", pdf_bytes, "application/pdf")
    pdf_stream = io.BytesIO(pdf_bytes)

    return StreamingResponse(
        pdf_stream,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=curriculum.pdf"}
    )


@router.post("/download/docx")
async def download_curriculum_docx(
    curriculum_data: dict = Body(...),
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=404, detail="찾을 수 없습니다")
    
    cur_title = curriculum_data.get("cur_title", "커리큘럼")
    week_plan = curriculum_data.get("cur_week_plan", [])
    if isinstance(week_plan, str):
        week_plan = json.loads(week_plan)

    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    doc.add_heading("ArtiCulum 과제 템플릿 패키지", 0)
    doc.add_paragraph(f"{cur_title} 실무 온보딩 과정").bold = True

    # 1. 과정 개요
    doc.add_heading("1. 과정 개요", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '과정명'
    hdr_cells[1].text = cur_title

    row1_cells = table.rows[1].cells
    row1_cells[0].text = '대상 직무'
    row1_cells[1].text = curriculum_data.get("cur_target_job", "미지정")

    row2_cells = table.rows[2].cells
    row2_cells[0].text = '교육 기간'
    row2_cells[1].text = f"{curriculum_data.get('cur_duration_weeks', 0)}주"

    row3_cells = table.rows[3].cells
    row3_cells[0].text = '교육 목적'
    row3_cells[1].text = curriculum_data.get("cur_learning_goal", "미지정")

    doc.add_paragraph("\n")

    # 2. 전체 과제 흐름
    doc.add_heading("2. 전체 과제 흐름", level=1)
    flow_table = doc.add_table(rows=1, cols=3)
    flow_table.style = 'Table Grid'
    flow_hdr = flow_table.rows[0].cells
    flow_hdr[0].text = '주차'
    flow_hdr[1].text = '주제'
    flow_hdr[2].text = '핵심 산출물'

    for week in week_plan:
        row = flow_table.add_row().cells
        row[0].text = f"{week.get('week')}주차"
        row[1].text = str(week.get('theme', ''))
        tasks = ", ".join([a.get('title', '') for a in week.get('assignments', [])])
        row[2].text = tasks
        
    doc.add_paragraph("\n")

    # 3. 주차별 과제 템플릿
    doc.add_heading("3. 주차별 과제 템플릿", level=1)
    for week in week_plan:
        doc.add_heading(f"{week.get('week')}주차 과제 템플릿", level=2)
        doc.add_paragraph(f"주제: {week.get('theme')}")
        doc.add_paragraph(f"학습 목표: {week.get('learning_objective')}")
        
        for idx, a in enumerate(week.get('assignments', [])):
            doc.add_heading(f"과제명: {a.get('title')}", level=3)
            doc.add_paragraph(f"제출 형태: {a.get('expected_output_format')}")
            
            p_guide = doc.add_paragraph("수행 방법:\n")
            for step in a.get('step_by_step_guide', []):
                p_guide.add_run(f"  - {step}\n")
                
        ig = week.get('instructor_guide', {})
        if ig and ig.get('check_points'):
            doc.add_heading("자기 점검 체크리스트", level=3)
            for cp in ig.get('check_points', []):
                doc.add_paragraph(f"☐ {cp}")
                
        if ig and ig.get('coaching_questions'):
            doc.add_heading("사수 피드백 질문", level=3)
            for idx, cq in enumerate(ig.get('coaching_questions', [])):
                doc.add_paragraph(f"{idx+1}. {cq}")
        
        doc.add_paragraph("\n")

    # 4. 공통 평가 루브릭
    doc.add_heading("4. 공통 평가 루브릭", level=1)
    rubric_table = doc.add_table(rows=5, cols=4)
    rubric_table.style = 'Table Grid'
    
    r_hdr = rubric_table.rows[0].cells
    r_hdr[0].text = '평가 항목'
    r_hdr[1].text = '우수'
    r_hdr[2].text = '보통'
    r_hdr[3].text = '미흡'
    
    rubric_data = [
        ("구체성", "행동, 수량, 순서가 명확하다", "대체로 구체적이나 일부 모호하다", "추상적 표현이 많다"),
        ("현장 적용성", "실제 업무에 바로 사용할 수 있다", "일부 수정 후 사용할 수 있다", "현장 상황과 연결이 약하다"),
        ("근거 제시", "명확한 데이터나 상황적 근거가 있다", "근거가 일부만 포함되어 있다", "개인 느낌 위주로 작성되었다"),
        ("연결성", "이전 주차 산출물을 다음 과제에 활용했다", "일부만 활용했다", "전혀 활용하지 않았다"),
    ]
    
    for i, row_data in enumerate(rubric_data, start=1):
        cells = rubric_table.rows[i].cells
        cells[0].text = row_data[0]
        cells[1].text = row_data[1]
        cells[2].text = row_data[2]
        cells[3].text = row_data[3]

    file_stream = io.BytesIO()
    doc.save(file_stream)
    docx_bytes = file_stream.getvalue()
    _save_curriculum_export(
        curriculum_data,
        "docx",
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=curriculum.docx"}
    )
    

@router.post("/generate-template", response_model=TemplateGenerateResponse)
@limiter.limit("20/hour")
def generate_template_api(
    request: Request,
    body: TemplateGenerateRequest,
    current_user: User = Depends(get_current_user),
):
    if current_user.user_role not in {"m", "a"}:
        raise HTTPException(status_code=403, detail="Only manager/admin can generate templates")

    try:
        val_result = curriculum_service.validate_curriculum_input(
            cur_title=body.assignment_title, 
            cur_target_job=body.theme or "",
            cur_target_industry="", 
            cur_learning_goal=body.learning_objective or "",
            required_content=" ".join(body.step_by_step_guide) if body.step_by_step_guide else ""
        )
        if not val_result.is_valid:
            raise HTTPException(status_code=400, detail=val_result.reason)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"템플릿 입력값 검증 실패: {exc}")

    try:
        html_content = curriculum_service.generate_assignment_template(
            theme=body.theme or "",
            learning_objective=body.learning_objective or "",
            assignment_title=body.assignment_title,
            step_by_step_guide=body.step_by_step_guide,
            expected_output_format=body.expected_output_format or "지정되지 않음",
        )
        return TemplateGenerateResponse(template_content=html_content)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"템플릿 AI 생성 실패: {exc}")
    

@router.delete("/{cur_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_curriculum(
    cur_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    curriculum = (
        db.query(Curriculum)
        .filter(Curriculum.cur_id == cur_id, Curriculum.cur_deleted_at.is_(None))
        .first()
    )
    if not curriculum or (current_user.user_role != "a" and curriculum.cur_creator_id != current_user.user_id):
        raise HTTPException(status_code=404, detail="커리큘럼을 찾을 수 없습니다")

    KST = timezone(timedelta(hours=9))
    curriculum.cur_deleted_at = datetime.now(KST)
    db.commit()
